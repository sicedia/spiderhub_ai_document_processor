import os
import logging
import time
import tempfile
from typing import Dict, Any, Optional
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from src.documentReport import DocumentReport

logger = logging.getLogger(__name__)

class WordTemplateGenerator:
    """Generate Word documents from DocumentReport using templates."""
    
    def __init__(self, template_path: Optional[str] = None):
        """
        Initialize the template generator.
        
        Args:
            template_path: Path to the Word template file (.docx)
        """
        self.template_path = template_path
        
    def load_template(self, template_path: str) -> Document:
        """
        Load a Word template document.
        
        Args:
            template_path: Path to the template file
            
        Returns:
            Document object loaded from template
        """
        if os.path.exists(template_path):
            logger.info(f"Loading template from: {template_path}")
            return Document(template_path)
        else:
            logger.warning(f"Template not found at {template_path}, creating blank document")
            return Document()
    
    def _safe_save(self, doc: Document, output_path: str, max_retries: int = 3):
        """
        Safely save document with retry mechanism.
        
        Args:
            doc: Document object to save
            output_path: Path to save the document
            max_retries: Maximum number of retry attempts
        """
        for attempt in range(max_retries):
            try:
                # Check if file exists and try to remove it first
                if os.path.exists(output_path):
                    try:
                        os.remove(output_path)
                        logger.info(f"Removed existing file: {output_path}")
                    except PermissionError:
                        # If we can't remove, try with a temporary name
                        timestamp = int(time.time())
                        base, ext = os.path.splitext(output_path)
                        output_path = f"{base}_{timestamp}{ext}"
                        logger.warning(f"File in use, saving as: {output_path}")
                
                # Save the document
                doc.save(output_path)
                logger.info(f"Document saved successfully to: {output_path}")
                return output_path
                
            except PermissionError as e:
                if attempt < max_retries - 1:
                    logger.warning(f"Permission denied (attempt {attempt + 1}), retrying in 2 seconds...")
                    time.sleep(2)
                    
                    # Try with timestamp suffix
                    timestamp = int(time.time())
                    base, ext = os.path.splitext(output_path)
                    output_path = f"{base}_{timestamp}{ext}"
                else:
                    logger.error(f"Failed to save after {max_retries} attempts: {e}")
                    raise
            except Exception as e:
                logger.error(f"Unexpected error saving document: {e}")
                raise
        
        return output_path
    
    def _get_available_styles(self, doc: Document) -> Dict[str, list]:
        """Get available styles from the template document."""
        paragraph_styles = []
        table_styles = []
        
        try:
            for style in doc.styles:
                if style.type == 1:  # Paragraph styles
                    paragraph_styles.append(style.name)
                elif style.type == 3:  # Table styles
                    table_styles.append(style.name)
        except Exception as e:
            logger.warning(f"Error getting styles: {e}")
        
        return {
            'paragraph': paragraph_styles,
            'table': table_styles
        }
    
    def replace_placeholder(self, doc: Document, placeholder: str, replacement: str):
        """
        Replace placeholder text in the document while preserving formatting.
        """
        # Replace in paragraphs
        for paragraph in doc.paragraphs:
            if placeholder in paragraph.text:
                # Preserve the original style
                original_style = paragraph.style
                paragraph.text = paragraph.text.replace(placeholder, replacement)
                paragraph.style = original_style
    
        # Replace in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if placeholder in paragraph.text:
                            original_style = paragraph.style
                            paragraph.text = paragraph.text.replace(placeholder, replacement)
                            paragraph.style = original_style
    
        # Replace in headers and footers
        for section in doc.sections:
            # Header
            header = section.header
            for paragraph in header.paragraphs:
                if placeholder in paragraph.text:
                    original_style = paragraph.style
                    paragraph.text = paragraph.text.replace(placeholder, replacement)
                    paragraph.style = original_style
            
            # Footer
            footer = section.footer
            for paragraph in footer.paragraphs:
                if placeholder in paragraph.text:
                    original_style = paragraph.style
                    paragraph.text = paragraph.text.replace(placeholder, replacement)
                    paragraph.style = original_style
    
    def add_bullet_list(self, doc: Document, items: list, heading: str = None):
        """
        Add a bullet list to the document.
        
        Args:
            doc: Document object
            items: List of items to add
            heading: Optional heading for the list
        """
        if heading:
            doc.add_heading(heading, level=2)
        
        if not items:
            p = doc.add_paragraph("No items available.")
            return
            
        for item in items:
            try:
                doc.add_paragraph(item, style='List Bullet')
            except:
                # Fallback if List Bullet style doesn't exist
                p = doc.add_paragraph(item)
                p.style = 'Normal'
    
    def add_table(self, doc: Document, data: Dict[str, list], heading: str = None):
        """
        Add a table to the document.
        
        Args:
            doc: Document object
            data: Dictionary with categories and items
            heading: Optional heading for the table
        """
        if heading:
            doc.add_heading(heading, level=2)
        
        if not data:
            doc.add_paragraph("No data available.")
            return
        
        # Get available table styles
        available_styles = self._get_available_styles(doc)
        table_style = 'Table Grid' if 'Table Grid' in available_styles['table'] else None
        
        # Create table
        table = doc.add_table(rows=1, cols=2)
        if table_style:
            table.style = table_style
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        
        # Add header row
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Category'
        hdr_cells[1].text = 'Items'
        
        # Make header bold
        for cell in hdr_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
        
        # Add data rows
        for category, items in data.items():
            row_cells = table.add_row().cells
            row_cells[0].text = category
            row_cells[1].text = "; ".join(items) if isinstance(items, list) else str(items)
    
    def create_from_template(self, report: DocumentReport, template_path: str, output_path: str):
        """
        Create a Word document from a template using the report data.
        
        Args:
            report: DocumentReport object with data
            template_path: Path to the Word template
            output_path: Path to save the generated document
        """
        try:
            # Load template
            doc = self.load_template(template_path)
            
            # Replace basic placeholders
            placeholders = {
                "{{date}}": report.date or "Date not specified",
                "{{location}}": report.location or "Location not specified",
                "{{executive_summary}}": report.executive_summary or "Executive summary not available.",
            }
            
            for placeholder, replacement in placeholders.items():
                self.replace_placeholder(doc, placeholder, replacement)
            
            # Find and replace list placeholders
            self._replace_list_placeholders(doc, report)
            
            # Find and replace table placeholders
            self._replace_table_placeholders(doc, report)
            
            # Save document with retry mechanism
            final_path = self._safe_save(doc, output_path)
            return final_path
            
        except Exception as e:
            logger.error(f"Error creating document from template: {e}")
            raise
    
    def _replace_list_placeholders(self, doc: Document, report: DocumentReport):
        """Replace list placeholders in the document."""
        list_placeholders = {
            "{{characteristics_list}}": report.characteristics,
            "{{applications_list}}": report.practical_applications,
            "{{commitments_list}}": report.commitments
        }
        
        for placeholder, items in list_placeholders.items():
            self._insert_list_at_placeholder(doc, placeholder, items)
    
    def _replace_table_placeholders(self, doc: Document, report: DocumentReport):
        """Replace table placeholders in the document."""
        table_placeholders = {
            "{{themes_table}}": report.themes,
            "{{actors_table}}": report.actors
        }
        
        for placeholder, data in table_placeholders.items():
            self._insert_table_at_placeholder(doc, placeholder, data)
    
    def _insert_list_at_placeholder(self, doc: Document, placeholder: str, items: list):
        """Insert a bullet list at the placeholder location using template styles."""
        paragraphs = list(doc.paragraphs)
        
        for paragraph in paragraphs:
            if placeholder in paragraph.text:
                # Capturar el estilo del párrafo que contiene el placeholder
                original_style = paragraph.style
                original_alignment = paragraph.alignment
                
                # Remove the placeholder
                paragraph.text = paragraph.text.replace(placeholder, "")
                
                if items:
                    # Buscar estilos de lista disponibles en la plantilla
                    available_styles = self._get_available_styles(doc)
                    
                    # Priorizar estilos de la plantilla
                    list_style = None
                    for style_name in available_styles['paragraph']:
                        if 'list' in style_name.lower() or 'bullet' in style_name.lower():
                            list_style = style_name
                            break
                    
                    # Fallback a 'List Bullet' o estilo original
                    if not list_style:
                        list_style = 'List Bullet' if 'List Bullet' in available_styles['paragraph'] else original_style
                    
                    # Insert items preservando formato
                    for item in reversed(items):
                        try:
                            p = doc.add_paragraph(item, style=list_style)
                            # Preservar alineación si es posible
                            if original_alignment:
                                p.alignment = original_alignment
                        except:
                            # Fallback completo
                            p = doc.add_paragraph(item)
                            p.style = original_style
                            if original_alignment:
                                p.alignment = original_alignment
                        
                        paragraph._element.addnext(p._element)
                else:
                    p = doc.add_paragraph("No items available.")
                    p.style = original_style
                    if original_alignment:
                        p.alignment = original_alignment
                    paragraph._element.addnext(p._element)
                break
    
    def _insert_table_at_placeholder(self, doc: Document, placeholder: str, data: Dict[str, list]):
        """Insert a table at the placeholder location using template styles."""
        paragraphs = list(doc.paragraphs)
        
        for paragraph in paragraphs:
            if placeholder in paragraph.text:
                # Remove the placeholder
                paragraph.text = paragraph.text.replace(placeholder, "")
                
                if data:
                    # Buscar tabla de referencia en la plantilla
                    reference_table_style = None
                    reference_font = None
                    
                    # Si hay tablas existentes, usar su estilo
                    if doc.tables:
                        reference_table = doc.tables[0]
                        reference_table_style = reference_table.style
                        
                        # Capturar formato de fuente de la primera celda
                        try:
                            first_cell = reference_table.rows[0].cells[0]
                            for para in first_cell.paragraphs:
                                for run in para.runs:
                                    if run.font.name or run.font.size:
                                        reference_font = run.font
                                        break
                                if reference_font:
                                    break
                        except:
                            pass
                    
                    # Crear tabla con estilo de referencia
                    table = doc.add_table(rows=1, cols=2)
                    
                    if reference_table_style:
                        table.style = reference_table_style
                    else:
                        # Fallback a estilos disponibles
                        available_styles = self._get_available_styles(doc)
                        table_style = 'Table Grid' if 'Table Grid' in available_styles['table'] else None
                        if table_style:
                            table.style = table_style
                    
                    table.alignment = WD_TABLE_ALIGNMENT.LEFT
                    
                    # Set column widths
                    try:
                        table.columns[0].width = Inches(2.0)
                        table.columns[1].width = Inches(4.0)
                    except:
                        pass
                    
                    # Add header con formato de referencia
                    hdr_cells = table.rows[0].cells
                    hdr_cells[0].text = 'Category'
                    hdr_cells[1].text = 'Items'
                    
                    # Aplicar formato de la plantilla
                    for cell in hdr_cells:
                        for paragraph_cell in cell.paragraphs:
                            for run in paragraph_cell.runs:
                                run.font.bold = True
                                if reference_font:
                                    if reference_font.name:
                                        run.font.name = reference_font.name
                                    if reference_font.size:
                                        run.font.size = reference_font.size
                    
                    # Add data con formato consistente
                    for category, items in data.items():
                        row_cells = table.add_row().cells
                        row_cells[0].text = category
                        row_cells[1].text = "; ".join(items) if isinstance(items, list) else str(items)
                        
                        # Aplicar formato de referencia a datos
                        if reference_font:
                            for cell in row_cells:
                                for paragraph_cell in cell.paragraphs:
                                    for run in paragraph_cell.runs:
                                        if reference_font.name:
                                            run.font.name = reference_font.name
                                        if reference_font.size:
                                            run.font.size = reference_font.size
                
                paragraph._element.addnext(table._element)
            else:
                p = doc.add_paragraph("No data available.")
                paragraph._element.addnext(p._element)
            break
    
    def create_default_document(self, report: DocumentReport, output_path: str):
        """
        Create a Word document with default formatting when no template is provided.
        
        Args:
            report: DocumentReport object with data
            output_path: Path to save the generated document
        """
        try:
            doc = Document()
            
            # Title
            title = doc.add_heading(report.title or "Document Analysis Report", 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Document info
            doc.add_heading('Document Information', level=1)
            info_table = doc.add_table(rows=3, cols=2)
            info_table.style = 'Table Grid'
            
            info_data = [
                ('Date', report.date or "Not specified"),
                ('Location', report.location or "Not specified"),
                ('Score', str(report.score) if report.score else "Not calculated")
            ]
            
            for i, (label, value) in enumerate(info_data):
                row = info_table.rows[i]
                row.cells[0].text = label
                row.cells[1].text = value
            
            # Executive Summary
            doc.add_heading('Executive Summary', level=1)
            doc.add_paragraph(report.executive_summary or "Executive summary not available.")
            
            # Characteristics
            self.add_bullet_list(doc, report.characteristics, 'Key Characteristics')
            
            # Themes
            self.add_table(doc, report.themes, 'Main Themes')
            
            # Actors
            self.add_table(doc, report.actors, 'Actors and Stakeholders')
            
            # Applications
            self.add_bullet_list(doc, report.practical_applications, 'Practical Applications')
            
            # Commitments
            self.add_bullet_list(doc, report.commitments, 'Commitments')
            
            # Extra data if available
            if report.extra_data:
                doc.add_heading('Additional Information', level=1)
                for key, value in report.extra_data.items():
                    doc.add_paragraph(f"{key}: {value}")
            
            # Save document with retry mechanism
            final_path = self._safe_save(doc, output_path)
            return final_path
            
        except Exception as e:
            logger.error(f"Error creating default document: {e}")
            raise

def generate_word_from_template(report: DocumentReport, template_path: str = None, 
                               output_path: str = "reports", filename_base: str = "report") -> str:
    """
    Generate a Word document from a DocumentReport using a template.
    
    Args:
        report: DocumentReport object with extracted data
        template_path: Path to Word template file (optional)
        output_path: Directory to save the output (optional)
        filename_base: Base filename without extension
        
    Returns:
        Path to the generated Word document
    """
    generator = WordTemplateGenerator(template_path)
    
    # Set default output path if not provided
    if not output_path:
        output_path = os.path.join(os.getcwd(), "reports")
    
    os.makedirs(output_path, exist_ok=True)
    
    # Clean filename to avoid Windows path issues
    clean_filename = "".join(c for c in filename_base if c.isalnum() or c in (' ', '-', '_')).rstrip()
    clean_filename = clean_filename.replace(' ', '_')
    
    # Full output file path
    output_file = os.path.join(output_path, f"{clean_filename}.docx")
    
    try:
        if template_path and os.path.exists(template_path):
            final_path = generator.create_from_template(report, template_path, output_file)
        else:
            final_path = generator.create_default_document(report, output_file)
        
        return final_path
        
    except Exception as e:
        logger.error(f"Error generating Word document: {e}")
        raise